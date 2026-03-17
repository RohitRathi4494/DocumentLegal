import json
from docx import Document

def generate_example():
    # 1. Create DOCX Template
    document = Document()
    document.add_heading('Rent Agreement', 0)
    
    document.add_paragraph('This Rent Agreement is made in {{agreement_city}} on {{agreement_date}}.')
    document.add_paragraph('Between Lessor: {{lessor_name}} residing at {{lessor_address}} (hereinafter referred to as the "Landlord")')
    document.add_paragraph('AND the following Tenants:')
    
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tenant Name'
    hdr_cells[1].text = 'Aadhar Number'
    hdr_cells[2].text = 'Permanent Address'
    
    # DocxTemplate loop syntax for tables
    row = table.add_row().cells
    row[0].text = '{% tr for tenant in tenants %}{{tenant.name}}'
    row[1].text = '{{tenant.aadhar}}'
    row[2].text = '{{tenant.address}}{% tr endfor %}'
    
    document.add_paragraph('\nProperty Details:')
    document.add_paragraph('The Landlord agrees to rent the property located at {{property_address}} comprising of {{property_type}}.')
    
    document.add_paragraph('\nFinancial Terms:')
    document.add_paragraph('1. Monthly Rent: Rs. {{monthly_rent}} ({{monthly_rent_words}} rupees only).')
    document.add_paragraph('2. Security Deposit: Rs. {{security_deposit}}.')
    
    document.add_paragraph('\nLease Terms:')
    document.add_paragraph('1. Lease Start Date: {{lease_start_date}}')
    document.add_paragraph('2. Lease Expiration Date: {{lease_end_date}} ({{lease_duration_months}} Months)')
    document.add_paragraph('3. Notice Period: {{notice_period_months}} Months')
    
    document.add_paragraph('\n{% if furniture_items %}')
    document.add_paragraph('Furniture & Fixtures Provided: {{furniture_items}}')
    document.add_paragraph('{% endif %}')
    
    document.add_paragraph('\n\n_______________________\nLessor Signature')
    
    document.save('../example_rent_agreement.docx')
    print("Expanded example_rent_agreement.docx generated.")

if __name__ == '__main__':
    generate_example()
