import docx
import re

def extract_variables(docx_path):
    doc = docx.Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    text = "\n".join(full_text)
    # Match {{ variable_name }}
    variables = re.findall(r'\{\{\s*(.*?)\s*\}\}', text)
    return sorted(list(set(variables)))

if __name__ == "__main__":
    path = "d:/MYDOCWRITER/backend/media/templates/Rent_Agreement_Template.docx"
    try:
        vars = extract_variables(path)
        print(f"Variables found in {path}:")
        for v in vars:
            print(f"- {v}")
    except Exception as e:
        print(f"Error: {e}")
